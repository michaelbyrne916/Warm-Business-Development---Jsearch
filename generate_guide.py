from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)

# ── Helpers ───────────────────────────────────────────────────
def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)   # dark navy
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # navy fill
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3564')
        tcPr.append(shd)
    # data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    doc.add_paragraph()   # spacing after table

# ══════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════
title = doc.add_heading('10 Pillars Solutions', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x35, 0x64)

sub = doc.add_heading('BD Automation Workflow — User Guide', 2)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# OVERVIEW
# ══════════════════════════════════════════════════════════════
heading('Overview')
body(
    'This system automatically finds companies actively hiring in your target markets, '
    'identifies the right contact, drafts personalized outreach emails, and manages the '
    'full send sequence. Your daily involvement is limited to two tasks: setting targets '
    'and approving emails.'
)

# ══════════════════════════════════════════════════════════════
# THE 5 WORKFLOWS
# ══════════════════════════════════════════════════════════════
heading('The 5 Workflows')
add_table(
    ['Workflow', 'What It Does', 'When It Runs'],
    [
        ['WF1: Lead Discovery',      'Finds companies actively hiring for your target roles',             'Weekdays 7am UTC'],
        ['WF2: Contact Discovery',   'Finds the best contact at each company via Hunter.io',              'Triggered by WF1'],
        ['WF3: Outreach Draft',      'Writes a 3-email sequence per contact using GPT-4o',               'Triggered by WF2'],
        ['WF4: Approved Send',       'Sends approved emails and schedules follow-ups',                   'On demand / scheduled'],
        ['WF5: Maintenance & Dedupe','Archives completed/stale rows, deduplicates the queue',            'Daily 6am'],
    ]
)

# ══════════════════════════════════════════════════════════════
# STEP 1 — TARGETING
# ══════════════════════════════════════════════════════════════
heading('Step 1 — Set Your Targets')
body(
    'Open the Google Sheet and go to the Targeting tab. Each row is one search '
    'configuration. Set active = TRUE to enable a row, FALSE to pause it.'
)
add_table(
    ['Column', 'What to Enter', 'Example'],
    [
        ['active',                  'TRUE to enable, FALSE to pause',                         'TRUE'],
        ['search_id',               'Unique label for this search',                           'sac-gis-analyst'],
        ['role_keywords',           'Semicolon-separated job titles to search',               'GIS Analyst; GIS Specialist'],
        ['locations',               'Semicolon-separated cities/states',                      'Sacramento, CA; Roseville, CA'],
        ['employment_type',         'FULLTIME, PARTTIME, CONTRACT, or blank',                 'FULLTIME'],
        ['excluded_keywords',       'Words that disqualify a job if found in the posting',    'intern; remote; junior'],
        ['excluded_company_types',  'Company types to skip',                                  'staffing; consulting'],
    ]
)
body('You can have multiple active rows — one per market or role category you are targeting.')

# ══════════════════════════════════════════════════════════════
# STEP 2 — AUTOMATIC MORNING PIPELINE
# ══════════════════════════════════════════════════════════════
heading('Step 2 — What Happens Automatically Each Weekday Morning')
body('You do not need to do anything for this part. The system runs in sequence:')
bullet('WF1 reads your Targeting tab and searches for matching job postings via JSearch.')
bullet('Jobs are scored (0–100). Staffing firms and large enterprises (>1,000 employees) are automatically excluded.')
bullet('Qualified companies are written to the Opportunities tab.')
bullet('WF2 immediately runs for each qualified company — enriches company data via Clearbit and finds the best contact via Hunter.io (confidence score required ≥ 70).')
bullet('Contacts are written to the Contacts tab.')
bullet('WF3 immediately runs — generates a personalized 3-email sequence (intro, follow-up, final touch) for each contact.')
bullet('All draft emails land in the Outreach Queue tab with approval_status = pending.')

# ══════════════════════════════════════════════════════════════
# STEP 3 — DAILY REVIEW
# ══════════════════════════════════════════════════════════════
heading('Step 3 — Your Daily Review (Outreach Queue Tab)')
body(
    'Open the Outreach Queue tab each morning. For each pending row, review the '
    'contact name, company, job title, and why_this_message_is_relevant. Read '
    'email_1_subject and email_1_body (the intro that will go out). You can also '
    'preview email_2 and email_3 (follow-ups at 3 and 7 days).'
)
body('Then set approval_status to one of the following values:')
add_table(
    ['Value', 'Meaning'],
    [
        ['approved',        'Send this email — WF4 will pick it up on the next run'],
        ['rejected',        'Skip this contact — WF5 will archive it tonight'],
        ['(leave pending)', 'Decide later — WF5 auto-archives after 14 days of inactivity'],
    ]
)

# ══════════════════════════════════════════════════════════════
# STEP 4 — SENDING
# ══════════════════════════════════════════════════════════════
heading('Step 4 — Sending Approved Emails (WF4)')
body(
    'Once you have set rows to approved, run WF4: Approved Send from the n8n editor '
    '(or ask to have it scheduled automatically).'
)
body('WF4 will:')
bullet('Send Email 1 via Gmail with your signature and unsubscribe line.')
bullet('Schedule Email 2 (follow-up) for 3 days later.')
bullet('Schedule Email 3 (final touch) for 7 days later.')
bullet('Update sent_status = sent and log the sent_timestamp.')
body('You do not need to do anything for the follow-ups — WF4 handles the full sequence automatically.')

# ══════════════════════════════════════════════════════════════
# STEP 5 — MAINTENANCE
# ══════════════════════════════════════════════════════════════
heading('Step 5 — Daily Cleanup (WF5, Automatic)')
body('Every morning at 6am, WF5 runs automatically and:')
bullet('Moves completed contacts (all 3 emails sent) to the Archive tab.')
bullet('Moves rejected rows to the Archive tab.')
bullet('Moves stale pending rows (pending + job posted > 14 days ago) to Archive.')
bullet('Removes duplicate contacts from the queue, keeping the most advanced stage.')
bullet('Rewrites only active rows back to the Outreach Queue.')
bullet('Emails you a daily summary with archived counts and breakdown by reason.')
body(
    'The Archive tab is your permanent history — nothing is ever deleted, only moved. '
    'The archived_reason column tells you exactly why each row was archived.'
)

# ══════════════════════════════════════════════════════════════
# SHEET REFERENCE
# ══════════════════════════════════════════════════════════════
heading('Reference: Google Sheet Tabs')
add_table(
    ['Tab', 'Purpose'],
    [
        ['Targeting',       'Your search configuration — roles, locations, and filters'],
        ['Opportunities',   'All job leads found (qualified, needs review, rejected)'],
        ['Contacts',        'All contacts discovered per opportunity'],
        ['Outreach Queue',  'Emails awaiting your approval and active send sequences'],
        ['Archive',         'Completed, rejected, and stale outreach — permanent history'],
        ['Errors',          'Workflow error log'],
    ]
)

# ══════════════════════════════════════════════════════════════
# DAILY ROUTINE SUMMARY
# ══════════════════════════════════════════════════════════════
heading('Daily Routine (2–5 minutes)')
bullet('Open the Outreach Queue tab in Google Sheets.')
bullet('For each pending row: set approval_status to approved or rejected.')
bullet('Run WF4 in n8n to send approved emails.')
bullet('Done — WF5 handles all cleanup automatically overnight.')

# ══════════════════════════════════════════════════════════════
# TIPS
# ══════════════════════════════════════════════════════════════
heading('Tips & Notes')
bullet('You can add multiple rows to the Targeting tab for different cities or role types simultaneously.')
bullet('If a contact is wrong, set approval_status = rejected — do not delete the row. WF5 will archive it.')
bullet('The Opportunities and Contacts tabs are source data — do not delete rows from them.')
bullet('WF5 runs before WF1 each day, so your queue is clean before new leads arrive.')
bullet('All emails are sent from your Gmail account (michael@10pillarssolutions.com) with your signature.')

# ══════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════
out = r'c:\Users\Michael\OneDrive\Documents\Claude - N8N\10Pillars_BD_Workflow_Guide.docx'
doc.save(out)
print(f'Saved: {out}')
